#!/usr/bin/env python3
"""
Simple Offer Letter Generator
Fills the offer template with candidate data directly (no profile config.json needed)
Usage: python generate_simple.py <candidate_json> [template_docx] [output_dir]
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from docx import Document


def replace_in_runs(paragraph, find_text, replace_text):
    """Replace text in runs while preserving formatting"""
    full_text = paragraph.text
    if find_text not in full_text:
        return False
    
    # Try run-level replacement first
    for run in paragraph.runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, str(replace_text))
            return True
    
    # Fall back: placeholder is split across runs - rebuild first run
    if paragraph.runs:
        new_text = full_text.replace(find_text, str(replace_text))
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
        return True
    
    return False


def fill_offer_letter(template_path, data, output_path):
    """Fill offer letter template with candidate data"""
    doc = Document(template_path)

    # Build replacements map
    replacements = {
        '{{Candidate Name}}': data.get('name', ''),
        '{{CANDIDATE_NAME}}': data.get('name', ''),
        '{{Interview Date}}': data.get('test_date', data.get('current_date', '')),
        '{{INTERVIEW_DATE}}': data.get('test_date', data.get('current_date', '')),
        '{{Job Title}}': data.get('position', ''),
        '{{JOB_TITLE}}': data.get('position', ''),
        '{{Joining Date}}': data.get('start_date', ''),
        '{{JOINING_DATE}}': data.get('start_date', ''),
        '{{Probation Monthly Salary}}': data.get('salary', ''),
        '{{MONTHLY_SALARY}}': data.get('salary', ''),
        '{{Probation Period Months}}': data.get('probation_period', '3'),
        '{{Probation period}}': data.get('probation_period', '3'),
        '{{Current Date}}': data.get('current_date', ''),
        '{{CURRENT_DATE}}': data.get('current_date', ''),
        '{{Acceptance Date}}': '',
        '{{Offer Validity Days}}': '7',
        '{{OFFER_EXPIRY_DAYS}}': '7',
        '{{ongoing_salary}}': data.get('ongoing_salary', data.get('salary', '')),
        '{{ONGOING_SALARY}}': data.get('ongoing_salary', data.get('salary', '')),
        '{{company}}': data.get('company', ''),
        '{{COMPANY}}': data.get('company', ''),
    }

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for find_text, replace_text in replacements.items():
            if find_text in paragraph.text:
                replace_in_runs(paragraph, find_text, str(replace_text))

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for find_text, replace_text in replacements.items():
                        if find_text in paragraph.text:
                            replace_in_runs(paragraph, find_text, str(replace_text))

    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_simple.py <candidate.json> [template.docx] [output_dir]")
        sys.exit(1)

    candidate_file = sys.argv[1]
    template_path = sys.argv[2] if len(sys.argv) > 2 else "templates/offer_template.docx"
    output_dir = sys.argv[3] if len(sys.argv) > 3 else "output"

    with open(candidate_file, 'r') as f:
        data = json.load(f)

    name_clean = data['name'].replace(' ', '_')
    out_dir = Path(output_dir)
    out_dir.mkdir(exist_ok=True, parents=True)
    output_path = out_dir / f"offer_letter_{name_clean}.docx"

    fill_offer_letter(template_path, data, output_path)
    print(f"\n🎉 Offer letter ready for {data['name']}")
    print(f"   Output: {output_path}")


if __name__ == "__main__":
    main()
