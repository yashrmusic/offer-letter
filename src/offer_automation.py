#!/usr/bin/env python3
"""
Complete Offer Letter Automation
- Fill template
- Convert to PDF
- Send via email
"""

import json
import sys
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from docx import Document
from pathlib import Path
import subprocess

def fill_offer_letter(template_path, data, output_path):
    """Fill offer letter template with candidate data"""
    doc = Document(template_path)
    
    # Map JSON data fields to template placeholders ({{Placeholder}} format)
    replacements = {
        '{{Candidate Name}}': data.get('name', ''),
        '{{Interview Date}}': data.get('test_date', ''),  # Using test_date as interview date
        '{{Job Title}}': data.get('position', ''),
        '{{Joining Date}}': data.get('start_date', ''),
        '{{Offer Validity Days}}': '2',  # Default value
        '{{Probation Monthly Salary}}': data.get('salary', ''),
        '{{Probation Period Months}}': '3',  # Default value
        '{{Acceptance Date}}': '',  # To be filled by candidate
    }
    
    for paragraph in doc.paragraphs:
        for find_text, replace_text in replacements.items():
            if find_text in paragraph.text:
                paragraph.text = paragraph.text.replace(find_text, str(replace_text))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for find_text, replace_text in replacements.items():
                    if find_text in cell.text:
                        cell.text = cell.text.replace(find_text, str(replace_text))
    
    # Fill Schedule 1 (Compensation table - usually the second table)
    if len(doc.tables) >= 2:
        schedule_table = doc.tables[1]  # Schedule 1 is typically the second table
        
        # Parse salary - handle ranges like "13000-15000" by taking the lower bound
        salary_str = data.get('salary', '0')
        salary_str = salary_str.replace('/-', '').replace(',', '').strip()
        
        if '-' in salary_str:
            # Handle range - take the first (lower) value
            parts = salary_str.split('-')
            try:
                salary = float(parts[0])
            except:
                salary = 0
        else:
            try:
                salary = float(salary_str)
            except:
                salary = 0
        
        # Only fill 4 rows with the same amount (no calculations needed)
        # All values are the same: Basic = Gross = Fixed = CTC
        salary_formatted = f"{salary:,.0f}"
        
        compensation_data = {
            'Basic': salary_formatted,
            'Gross Compensation': salary_formatted,
            'Fixed Compensation': salary_formatted,
            'Cost to Company Compensation': salary_formatted,
        }
        
        # Fill the amounts in the second column (only for specified rows)
        for row in schedule_table.rows[1:]:  # Skip header row
            if len(row.cells) >= 2:
                component = row.cells[0].text.strip()
                if component in compensation_data:
                    row.cells[1].text = compensation_data[component]
    
    doc.save(output_path)
    return output_path

def convert_to_pdf(docx_path, pdf_path):
    """Convert .docx to PDF using LibreOffice"""
    try:
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(pdf_path.parent),
            str(docx_path)
        ], check=True)
        print(f"✅ PDF created: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")
        return None

def send_email(pdf_path, recipient_email, sender_email, sender_password):
    """Send offer letter PDF via email"""
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = "Offer Letter - Please Review and Sign"
    
    body = """
Dear Candidate,

Please find attached your offer letter.

Kindly review and let us know if you have any questions.

Best regards
"""
    
    msg.attach(MIMEText(body, 'plain'))
    
    with open(pdf_path, 'rb') as f:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename={pdf_path.name}')
        msg.attach(part)
    
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        print(f"✅ Email sent to {recipient_email}")
        return True
    except Exception as e:
        print(f"❌ Email failed: {e}")
        return False

def main():
    if len(sys.argv) < 2:
        print("Usage: ./offer_automation.py <data.json> [--pdf] [--email sender@email.com password]")
        sys.exit(1)
    
    template = Path("/Users/yashrakhiani/.gemini/antigravity/scratch/offer_template.docx")
    
    with open(sys.argv[1], 'r') as f:
        data = json.load(f)
    
    name_clean = data['name'].replace(' ', '_')
    docx_out = Path(f"/Users/yashrakhiani/.gemini/antigravity/scratch/offer_letter_{name_clean}.docx")
    
    print("📝 Filling offer letter...")
    fill_offer_letter(template, data, docx_out)
    print(f"✅ Created: {docx_out}")
    
    if '--pdf' in sys.argv:
        pdf_out = docx_out.with_suffix('.pdf')
        print("\n📄 Converting to PDF...")
        convert_to_pdf(docx_out, pdf_out)
    
    if '--email' in sys.argv:
        idx = sys.argv.index('--email')
        sender = sys.argv[idx + 1]
        password = sys.argv[idx + 2]
        
        pdf_file = docx_out.with_suffix('.pdf')
        if pdf_file.exists():
            print(f"\n📧 Sending email to {data['email']}...")
            send_email(pdf_file, data['email'], sender, password)
        else:
            print("❌ PDF not found. Generate PDF first with --pdf")

if __name__ == "__main__":
    main()
