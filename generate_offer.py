#!/usr/bin/env python3
"""
Format-preserving Offer Letter Generation
Uses simple text replacement within runs to preserve formatting
"""

import json
import sys
from pathlib import Path
from docx import Document
import subprocess


def load_profile(profile_name):
    """Load company profile configuration"""
    profile_path = Path('profiles') / profile_name / 'config.json'
    
    if not profile_path.exists():
        print(f"❌ Profile '{profile_name}' not found!")
        print(f"Available profiles: melange, urbanmistrii, decoarte")
        sys.exit(1)
    
    with open(profile_path, 'r') as f:
        return json.load(f)


def replace_in_runs(paragraph, find_text, replace_text):
    """Replace text in runs while preserving formatting"""
    for run in paragraph.runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text)
            return True
    return False


def fill_offer_letter(template_path, data, profile, output_path):
    """Fill offer letter template with candidate data - preserving formatting"""
    doc = Document(template_path)
    
    # Map JSON data fields to template placeholders
    replacements = {
        '{{Candidate Name}}': data.get('name', ''),
        '{{CANDIDATE_NAME}}': data.get('name', ''),
        '{{Interview Date}}': data.get('test_date', ''),
        '{{INTERVIEW_DATE}}': data.get('test_date', ''),
        '{{Job Title}}': data.get('position', ''),
        '{{JOB_TITLE}}': data.get('position', ''),
        '{{Joining Date}}': data.get('start_date', ''),
        '{{JOINING_DATE}}': data.get('start_date', ''),
        '{{Offer Validity Days}}': str(profile['offer_validity_days']),
        '{{Probation Monthly Salary}}': data.get('salary', ''),
        '{{MONTHLY_SALARY}}': data.get('salary', ''),
        '{{Probation Period Months}}': data.get('probation_period', str(profile['probation_months'])),
        '{{Probation period}}': data.get('probation_period', str(profile['probation_months'])),
        '{{Acceptance Date}}': '',
        '{{Current Date}}': data.get('current_date', ''),
        '{{CURRENT_DATE}}': data.get('current_date', ''),
        '{{OFFER_EXPIRY_DAYS}}': str(profile['offer_validity_days']),
    }
    
    # Replace in paragraphs - try run-level first, then fall back to paragraph-level
    for paragraph in doc.paragraphs:
        for find_text, replace_text in replacements.items():
            if find_text in paragraph.text:
                # Try run-level replacement first (preserves formatting)
                replaced = replace_in_runs(paragraph, find_text, str(replace_text))
                if not replaced:
                    # Fall back to rebuilding runs if placeholder spans multiple runs
                    inline = paragraph.runs
                    for i, run in enumerate(inline):
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, str(replace_text))
    
    # Replace in tables - need special handling for split placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for find_text, replace_text in replacements.items():
                        if find_text in paragraph.text:
                            # Try run-level first
                            replaced = replace_in_runs(paragraph, find_text, str(replace_text))
                            if not replaced and paragraph.runs:
                                # Placeholder is split across runs - rebuild
                                full_text = paragraph.text
                                new_text = full_text.replace(find_text, str(replace_text))
                                # Keep first run's formatting, put all text there
                                paragraph.runs[0].text = new_text
                                for run in paragraph.runs[1:]:
                                    run.text = ""
    
    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(pdf_path.parent),
            str(docx_path)
        ], check=True)
        return pdf_path
    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")
        return None


def generate_offer_internal(profile_name, data):
    """Programmatic interface for offer generation"""
    profile = load_profile(profile_name)
    
    # Prepare paths
    name_clean = data['name'].replace(' ', '_')
    output_dir = Path('output') / profile_name
    output_dir.mkdir(exist_ok=True, parents=True)
    
    docx_out = output_dir / f'offer_letter_{name_clean}.docx'
    template_path = Path(profile['template_docx'])
    
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    fill_offer_letter(template_path, data, profile, docx_out)
    
    # Convert to PDF
    pdf_out = docx_out.with_suffix('.pdf')
    convert_to_pdf(docx_out, pdf_out)
    
    return {
        'docx': docx_out,
        'pdf': pdf_out
    }


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 generate_offer.py <profile> <candidate.json>")
        print("\nProfiles:")
        print("  melange      - The Melange Studio")
        print("  urbanmistrii - Urban Mistrii")
        print("  decoarte     - Deco Arte")
        print("\nExample:")
        print("  python3 generate_offer.py melange examples/sample_candidate.json")
        sys.exit(1)
    
    profile_name = sys.argv[1]
    candidate_file = sys.argv[2]
    
    # Load profile and candidate data
    print(f"📋 Loading profile: {profile_name}")
    profile = load_profile(profile_name)
    
    print(f"👤 Loading candidate data: {candidate_file}")
    with open(candidate_file, 'r') as f:
        data = json.load(f)
    
    # Prepare paths
    name_clean = data['name'].replace(' ', '_')
    output_dir = Path('output') / profile_name
    output_dir.mkdir(exist_ok=True, parents=True)
    
    docx_out = output_dir / f'offer_letter_{name_clean}.docx'
    
    # Generate offer letter
    print(f"📝 Filling offer letter for {profile['company_name']}...")
    template_path = Path(profile['template_docx'])
    
    if not template_path.exists():
        print(f"❌ Template not found: {template_path}")
        print(f"💡 Please add the template to: {template_path}")
        sys.exit(1)
    
    fill_offer_letter(template_path, data, profile, docx_out)
    print(f"✅ Created: {docx_out}")
    
    # Convert to PDF
    pdf_out = docx_out.with_suffix('.pdf')
    print(f"\n📄 Converting to PDF...")
    convert_to_pdf(docx_out, pdf_out)
    print(f"✅ PDF created: {pdf_out}")
    
    print(f"\n🎉 Success! Offer letter ready for {data['name']}")
    print(f"   Company: {profile['company_name']}")
    print(f"   Output: {pdf_out}")

if __name__ == "__main__":
    main()
