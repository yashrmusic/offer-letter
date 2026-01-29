#!/usr/bin/env python3
"""
Web server for digital signature collection
Serves the signature interface and handles signature submission
"""

from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import json
import base64
from pathlib import Path
from PIL import Image
import io
from docx import Document
import google.generativeai as genai
import subprocess
from generate_offer import generate_offer_internal, load_profile as load_company_profile
from send_email import send_offer_email

app = Flask(__name__, 
            static_folder='web',
            template_folder='web')
CORS(app)

OUTPUT_DIR = Path('output')
WEB_DIR = Path('web')
EXAMPLES_DIR = Path('examples')

import os
from dotenv import load_dotenv

load_dotenv()

# Initialize Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-flash')
else:
    print("⚠️  Warning: GEMINI_API_KEY not found in environment variables.")
    model = None

@app.route('/')
def index():
    return send_file(WEB_DIR / 'index.html')

@app.route('/admin')
def admin():
    return send_file(WEB_DIR / 'admin.html')

@app.route('/api/profiles')
def get_profiles():
    """List available company profiles"""
    try:
        profiles_dir = Path('profiles')
        profiles = [d.name for d in profiles_dir.iterdir() if d.is_dir()]
        return jsonify({'success': True, 'profiles': profiles})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/generate-offer', methods=['POST'])
def generate_offer():
    """Generate offer letter from web form data"""
    try:
        data = request.json
        profile_name = data.get('profile')
        candidate_data = data.get('candidate')
        
        if not profile_name or not candidate_data:
            return jsonify({'success': False, 'message': 'Missing profile or candidate data'}), 400
            
        result = generate_offer_internal(profile_name, candidate_data)
        
        return jsonify({
            'success': True,
            'message': f"Offer generated for {candidate_data['name']}",
            'pdf_url': f"/api/offer-pdf/{candidate_data['name'].replace(' ', '_')}?profile={profile_name}"
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ai-parse', methods=['POST'])
def ai_parse():
    """Use Gemini to parse unstructured candidate info from a prompt"""
    try:
        data = request.json
        prompt = data.get('prompt')
        
        if not prompt:
            return jsonify({'success': False, 'message': 'No prompt provided'}), 400
            
        sys_instruction = """
        Extract candidate information from the following text and return it in JSON format.
        Fields to extract:
        - name: Full name
        - email: Email address
        - phone: Phone number
        - position: Job title/role
        - start_date: Joining date (e.g. 5 January, 2026)
        - salary: Monthly salary (numeric or range)
        - test_date: Interview date
        
        If a field is not found, leave it as an empty string. Output ONLY the JSON.
        """
        
        response = model.generate_content(f"{sys_instruction}\n\nCandidate Text: {prompt}")
        
        # Clean the response to ensure it's valid JSON
        text = response.text.strip()
        if text.startswith('```json'):
            text = text[7:]
        if text.endswith('```'):
            text = text[:-3]
        text = text.strip()
            
        parsed_data = json.loads(text)
        return jsonify({'success': True, 'data': parsed_data})
        
    except Exception as e:
        print(f"AI Error: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500
@app.route('/api/send-email', methods=['POST'])
def send_email():
    """Send generated offer letter via email"""
    try:
        data = request.json
        profile_name = data.get('profile')
        candidate_name = data.get('candidate_name')
        candidate_email = data.get('candidate_email') # Optional override
        
        if not profile_name or not candidate_name:
            return jsonify({'success': False, 'message': 'Missing profile or candidate name'}), 400
            
        profile = load_company_profile(profile_name)
        
        # Load candidate data to get correct name format and email
        name_clean = candidate_name.replace(' ', '_')
        pdf_path = OUTPUT_DIR / profile_name / f'offer_letter_{name_clean}.pdf'
        
        if not pdf_path.exists():
            return jsonify({'success': False, 'message': f"PDF not found at {pdf_path}. Generate it first."}), 404
            
        # We need the JSON file path for send_offer_email
        # For web-generated offers, we might not have a JSON file yet, or we can create a temporary one
        # Alternatively, refactor send_offer_email to take data directly
        
        # Let's check if the json exists in examples/
        candidate_file = EXAMPLES_DIR / f"{name_clean.lower()}.json"
        
        # If it doesn't exist, we'll create a minimal one for the email script
        if not candidate_file.exists():
            candidate_file = Path(f"/tmp/{name_clean.lower()}.json")
            with open(candidate_file, 'w') as f:
                json.dump({
                    "name": candidate_name,
                    "email": candidate_email,
                    "position": "Selected Position" # Fallback
                }, f)

        success = send_offer_email(str(candidate_file), pdf_path, profile, recipient_override=candidate_email)
        
        if success:
            return jsonify({'success': True, 'message': f"Email sent to {candidate_name}"})
        else:
            return jsonify({'success': False, 'message': "Failed to send email. Check logs."}), 500
            
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/offer-preview/<candidate_name>')
def get_offer_preview(candidate_name):
    """Get offer letter preview data"""
    try:
        profile_name = request.args.get('profile', 'melange')
        
        # Load candidate data
        # We try to find the candidate data in examples/ or in profile-specific dirs if they exist
        candidate_file = EXAMPLES_DIR / f'sample_candidate.json' # Default fallback
        
        # Optional: search for specific candidate json if needed
        # For now, we'll keep using the sample or look for one matching name
        potential_candidate_file = EXAMPLES_DIR / f'{candidate_name.lower()}.json'
        if potential_candidate_file.exists():
            candidate_file = potential_candidate_file
            
        with open(candidate_file, 'r') as f:
            data = json.load(f)
        
        return jsonify({
            'success': True,
            'candidate': {
                'name': data.get('name', ''),
                'position': data.get('position', ''),
                'start_date': data.get('start_date', ''),
                'interview_date': data.get('test_date', ''),
                'salary': data.get('salary', '')
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/offer-pdf/<candidate_name>')
def get_offer_pdf(candidate_name):
    """Serve the unsigned PDF"""
    try:
        profile_name = request.args.get('profile', 'melange')
        pdf_file = OUTPUT_DIR / profile_name / f'offer_letter_{candidate_name}.pdf'
        
        if not pdf_file.exists():
            # Try root output dir just in case
            pdf_file = OUTPUT_DIR / f'offer_letter_{candidate_name}.pdf'
            
        if pdf_file.exists():
            return send_file(pdf_file, mimetype='application/pdf')
        else:
            return jsonify({'success': False, 'message': f'PDF not found at {pdf_file}'}), 404
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/submit-signature', methods=['POST'])
def submit_signature():
    """Handle signature submission and create signed PDF"""
    try:
        data = request.json
        candidate_name = data.get('candidate', 'Mariya_Fatima')
        profile_name = data.get('profile', 'melange')
        signature_data = data.get('signature', '')
        signature_date = data.get('date', '')
        
        # Decode base64 signature image
        signature_base64 = signature_data.split(',')[1] if ',' in signature_data else signature_data
        signature_bytes = base64.b64decode(signature_base64)
        
        # Save signature image to output dir
        profile_output_dir = OUTPUT_DIR / profile_name
        profile_output_dir.mkdir(exist_ok=True, parents=True)
        
        sig_image_path = profile_output_dir / f'signature_{candidate_name}.png'
        with open(sig_image_path, 'wb') as f:
            f.write(signature_bytes)
        
        # Path to the unsigned DOCX
        docx_path = profile_output_dir / f'offer_letter_{candidate_name}.docx'
        
        if not docx_path.exists():
            # Try checking the root output dir
            docx_path = OUTPUT_DIR / f'offer_letter_{candidate_name}.docx'
            
        if not docx_path.exists():
             return jsonify({'success': False, 'message': f'Source DOCX not found at {docx_path}'}), 404
             
        signed_docx_path = profile_output_dir / f'offer_letter_{candidate_name}_signed.docx'
        
        # Copy and add signature to DOCX
        doc = Document(docx_path)
        
        # Find the signature placeholder or just add to the end
        found_placeholder = False
        for paragraph in doc.paragraphs:
            if 'Signature:' in paragraph.text or 'Candidate Signature' in paragraph.text:
                # Add some space
                p = doc.add_paragraph()
                # Add signature image
                run = p.add_run()
                run.add_picture(str(sig_image_path), width=Inches(2))
                # Add date
                doc.add_paragraph(f'Date: {signature_date}')
                found_placeholder = True
                break
        
        if not found_placeholder:
            # If no placeholder, just add at the end
            doc.add_paragraph('\nCandidate Signature:')
            doc.add_picture(str(sig_image_path), width=Inches(2))
            doc.add_paragraph(f'Date: {signature_date}')
            
        doc.save(signed_docx_path)
        
        # Convert to PDF
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(profile_output_dir),
            str(signed_docx_path)
        ], check=True)
        
        return jsonify({
            'success': True,
            'message': 'Signature submitted successfully',
            'signed_pdf_url': f'/api/signed-pdf/{candidate_name}?profile={profile_name}'
        })
        
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/signed-pdf/<candidate_name>')
def get_signed_pdf(candidate_name):
    """Serve the signed PDF"""
    try:
        profile_name = request.args.get('profile', 'melange')
        pdf_file = OUTPUT_DIR / profile_name / f'offer_letter_{candidate_name}_signed.pdf'
        
        if pdf_file.exists():
            return send_file(pdf_file, 
                           mimetype='application/pdf',
                           as_attachment=True,
                           download_name=f'offer_letter_{candidate_name}_signed.pdf')
        else:
            return jsonify({'success': False, 'message': 'Signed PDF not found'}), 404
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

if __name__ == '__main__':
    # Ensure output directory exists
    OUTPUT_DIR.mkdir(exist_ok=True)
    
    print("🚀 Starting signature collection server...")
    print("📝 Open http://localhost:5001 in your browser")
    print("✍️  Candidates can sign their offer letters digitally!")
    
    # Check if templates and output exist
    if not (OUTPUT_DIR / 'melange').exists():
        print("⚠️  Warning: output/melange directory not found. Have you generated any offers yet?")
        
    app.run(debug=True, port=5001)
