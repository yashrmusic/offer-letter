from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_template():
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Déco-Arte")
    run.bold = True
    run.font.size = Pt(24)
    
    subhead = doc.add_paragraph()
    subhead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subhead.add_run("The Bespoke Interior Architecture & Design Co")
    run.italic = True
    run.font.size = Pt(10)
    
    doc.add_paragraph() # Spacer
    
    # Date
    p = doc.add_paragraph("Date: {{Current Date}}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph() # Spacer
    
    # Title
    title = doc.add_paragraph("OFFER LETTER")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    title.runs[0].underline = True
    
    doc.add_paragraph() # Spacer
    
    # Candidate Name
    doc.add_paragraph("To,")
    doc.add_paragraph("{{Candidate Name}}")
    doc.add_paragraph()
    
    # Body
    body = doc.add_paragraph()
    body.add_run("Further to the personal interview held on ")
    body.add_run("{{Interview Date}}").bold = True
    body.add_run(", Deco Arte is pleased to offer you employment as ")
    body.add_run("{{Job Title}}").bold = True
    body.add_run(" starting ")
    body.add_run("{{Joining Date}}").bold = True
    body.add_run(". You will be stationed at our office at E-204, near Kailash Colony Metro Station, Block E, East of Kailash, New Delhi, Delhi 110065, India.")
    
    doc.add_paragraph()
    
    body2 = doc.add_paragraph()
    body2.add_run("Your salary during probation will be ")
    body2.add_run("INR {{Probation Monthly Salary}}").bold = True
    body2.add_run(" per month. The probation period will be for ")
    body2.add_run("{{Probation Period Months}} months").bold = True
    body2.add_run(". Please note that this offer is subject to obtaining satisfactory references from your past employer and university and completing other joining formalities, including signing the employment agreement and confidentiality agreement.")
    
    doc.add_paragraph()
    
    # Table
    doc.add_paragraph("Compensation Structure:", style='List Bullet')
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ["Fund", ""],
        ["Gratuity", "-"],
        ["Fixed Compensation", "INR {{Probation Monthly Salary}} / month"],
        ["Approx Value Of Benefits", "-"],
        ["Cost to Company (CTC)", "INR {{Probation Monthly Salary}} / month"]
    ]
    
    for i, (label, value) in enumerate(data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    
    # Notes
    notes = doc.add_paragraph()
    run = notes.add_run("Notes:")
    run.bold = True
    doc.add_paragraph("* Flexible pay encompasses allowances, including travel expenses for site visits and stationery utilized for office activities.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Validity
    validity = doc.add_paragraph()
    validity.add_run("This offer is valid for ")
    validity.add_run("{{Offer Validity Days}} days").bold = True
    validity.add_run(" from the date of issue.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph("Deco Arte")
    footer.runs[0].bold = True
    doc.add_paragraph("E-204, East Of Kailash, Delhi, 110048, India")
    doc.add_paragraph("T: [+91-9810281799]")
    doc.add_paragraph("E: hr@decoarte.com")
    doc.add_paragraph("www.deco-arte.in")
    
    # Save
    path = "/Users/yashrakhiani/Desktop/OfferLetterAutomation/profiles/decoarte/offer_template.docx"
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"Template created at {path}")

if __name__ == "__main__":
    create_template()
